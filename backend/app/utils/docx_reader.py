"""
DOCX文档读取工具 - 从Word文档中提取测试用例

支持格式:
- 段落中的序列 (以大写氨基酸字母为主的段落)
- 表格中的测试用例定义
- 特定标记的段落 (## 测试用例, **序列**:, etc.)
"""

from __future__ import annotations
from pathlib import Path
from typing import Optional


def read_docx_text(file_path: str) -> str:
    """读取docx文件的全部文本内容"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except ImportError:
        raise RuntimeError("python-docx未安装。请执行: pip install python-docx")


def extract_sequences_from_docx(file_path: str) -> List[dict]:
    """
    从DOCX文档中提取测试用的序列和目标

    识别逻辑:
    1. 查找连续的大写氨基酸序列 (>30个字符, 主要由ACDEFGHIKLMNPQRSTVWY组成)
    2. 序列前后的文本作为test_name和target
    3. 也支持表格格式的测试用例

    Returns:
        [{"sequence": str, "target": str, "test_name": str}, ...]
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx未安装。请执行: pip install python-docx")

    doc = Document(file_path)
    tests = []

    # AA字母集
    aa_set = set("ACDEFGHIKLMNPQRSTVWY")

    # 方法1: 从段落中提取序列
    all_text = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            all_text.append(text)

    current_name = ""
    current_target = "increase_strength"

    for text in all_text:
        # 检查是否是目标描述
        text_lower = text.lower()
        if any(kw in text_lower for kw in ["测试", "test", "用例", "目标", "target"]):
            current_name = text[:100]
            if "strength" in text_lower:
                current_target = "increase_strength"
            elif "elastic" in text_lower:
                current_target = "increase_elasticity"
            elif "toughness" in text_lower:
                current_target = "increase_toughness"
            elif "balanced" in text_lower or "综合" in text_lower:
                current_target = "balanced_optimization"
            continue

        # 检查是否是序列 (主要包含AA字母，长度>30)
        clean = text.replace(" ", "").replace("\n", "").replace("\r", "")
        if len(clean) >= 30:
            aa_ratio = sum(1 for c in clean if c in aa_set) / max(len(clean), 1)
            if aa_ratio > 0.85:
                tests.append({
                    "sequence": clean,
                    "target": current_target,
                    "test_name": current_name or f"序列_{len(tests)+1}",
                })
                current_name = ""

    # 方法2: 从表格中提取
    for table in doc.tables:
        headers = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if i == 0:
                headers = [h.lower() for h in cells]
                continue

            row_data = dict(zip(headers, cells))
            seq = None
            for key in ["sequence", "序列", "seq"]:
                if key in row_data:
                    seq = row_data[key].replace(" ", "").replace("\n", "")
                    break

            if seq and len(seq) >= 10:
                aa_ratio = sum(1 for c in seq if c in aa_set) / max(len(seq), 1)
                if aa_ratio > 0.85:
                    target = row_data.get("target", row_data.get("目标", "increase_strength"))
                    name = row_data.get("test_name", row_data.get("name", row_data.get("名称", f"表格测试_{len(tests)+1}")))
                    tests.append({"sequence": seq, "target": target, "test_name": name})

    return tests


def parse_docx_to_test_requests(file_path: str) -> List[dict]:
    """
    解析DOCX文档并转换为批量测试请求格式

    这是与 /api/v1/afp/batch-test 端点对接的便捷函数
    """
    return extract_sequences_from_docx(file_path)


def extract_text_and_tables(file_path: str) -> dict:
    """
    提取DOCX文档的完整结构和内容

    Returns:
        {"paragraphs": [str, ...], "tables": [[[str, ...], ...], ...]}
    """
    try:
        from docx import Document
        doc = Document(file_path)
        return {
            "paragraphs": [p.text for p in doc.paragraphs],
            "tables": [
                [[cell.text for cell in row.cells] for row in table.rows]
                for table in doc.tables
            ],
        }
    except ImportError:
        raise RuntimeError("python-docx未安装。请执行: pip install python-docx")
